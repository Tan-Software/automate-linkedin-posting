import datetime

from termcolor import colored
from docx import Document
from docx.shared import Inches
from pathlib import Path

def generate_initial_docx(input_date, input_hour):
    from termcolor import colored

    document = Document()

    document.add_heading('Post généré automatiquement par Tansoftware', 0)
    document.add_paragraph("🥳 Vous pouvez désormais éditer ce document afin qu'il puisse être posté automatiquement sur LinkedIn.")
    document.add_page_break()
    
    document_formated_name = str(input_date.replace("-", "_")) + "__" + str(input_hour.replace(":", "_"))
    
    scheduled_posts_dir = Path("./scheduled_posts")
    scheduled_posts_dir.mkdir(exist_ok=True)
    document_path = scheduled_posts_dir / f"___{document_formated_name}.docx"
    document.save(document_path)
    
    text = "Le document : " + document_path.name + ", a été généré dans le répertoire scheduled_posts"
    separator_aligned_to_text = "-" * (len(text) + 4)

    print(colored(f"\n{separator_aligned_to_text}", "yellow"))
    print(colored(f" {text}", "yellow"))
    print(colored(f"{separator_aligned_to_text}\n", "yellow"))
    print(colored("\nSi l'outil ne tourne pas en tâche de fond, pensez à le lancer avec la commande :", "blue"))
    print(colored("- python main.py", "red"))
    print(colored("\nPuis selectionner le choix : 2", "blue"))


def question_datetime(fmt, question, default):
    input_datetime = input(colored("\n " + question + " (au format " + default + ") ? ", 'green') \
        + colored("[Appuyez sur entrer pour mettre la valeur par défaut]", "blue") \
        + colored(" : ", "green")) or default

    try:
        datetime.datetime.strptime(input_datetime, fmt)

        return input_datetime
    except ValueError:
        error_msg = "\n Le format n'est pas valide, il doit être du type : " + default + "\n"
        print(colored(error_msg, "red"))
        return question_datetime(fmt, question, default)

def questions():
    input_date = question_datetime("%d-%m-%Y", "A quelle date voulez vous l'envoyer", datetime.datetime.now().strftime("%d-%m-%Y"))
    input_hour = question_datetime("%H:%M", "A quelle heure voulez vous l'envoyer", datetime.datetime.now().strftime("%H:%M"))

    generate_initial_docx(input_date, input_hour)

def init():
    questions()